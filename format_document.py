"""
BSU Engineering Portal — Document Formatter
============================================
Runs Pandoc to convert Markdown to .docx, then applies
professional academic formatting via python-docx.

Usage (in WSL):
    pip install python-docx --break-system-packages
    python3 format_document.py
"""

import subprocess
import sys
from pathlib import Path

# ─── CONFIGURATION ────────────────────────────────────────────────────────────

DOCS_DIR   = Path("/mnt/m/Graduation Project DevOps/1/BSU_Engineering_Portal_Graduation_Project-main")
MD_FILE    = Path("/mnt/c/Users/engmo/.gemini/antigravity/brain/b883b046-6920-47cc-9b6c-ba1158212ec1/section_1_cover_and_formal_pages.md")
HEADER_IMG = Path("/mnt/c/Users/engmo/.gemini/antigravity/brain/b883b046-6920-47cc-9b6c-ba1158212ec1/bsu_faculty_header_1774631573299.png")
RAW_DOCX   = DOCS_DIR / "Section1_raw.docx"
FINAL_DOCX = DOCS_DIR / "Section1_BSU_Engineering_Portal.docx"

# Colours (R, G, B)
NAVY  = (27,  58,  107)   # #1B3A6B — Heading 1
GOLD  = (180, 140,  30)   # #B48C1E — Heading 2
DARK  = (50,   50,  50)   # #323232 — Heading 3
LTBLU = (235, 241, 251)   # alternating table rows


# ─── STEP 1: PANDOC ───────────────────────────────────────────────────────────

def run_pandoc():
    print("▶  Running Pandoc …")
    cmd = [
        "pandoc", str(MD_FILE),
        "-o", str(RAW_DOCX),
        "--standalone",
        "--toc", "--toc-depth=3",
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print("✗  Pandoc failed:\n", result.stderr)
        sys.exit(1)
    print("✓  Pandoc done →", RAW_DOCX.name)


# ─── STEP 2: FORMATTING ───────────────────────────────────────────────────────

def apply_formatting():
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("✗  python-docx not found.")
        print("   Run: pip install python-docx --break-system-packages")
        sys.exit(1)

    print("▶  Applying professional formatting …")
    doc = Document(str(RAW_DOCX))

    # Print available styles so we can see what pandoc produced
    style_names = [s.name for s in doc.styles]
    print(f"   Styles found: {style_names[:20]}")

    # ── Margins: A4, Egyptian-thesis standard ─────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.0)
        sec.page_height   = Cm(29.7)
        sec.page_width    = Cm(21.0)

    # ── Safe style finder ─────────────────────────────────────────────────────
    def get_style(fragment):
        for s in doc.styles:
            if fragment.lower() in s.name.lower():
                return s
        return None

    # ── Global style objects (if they exist in the document) ──────────────────
    mappings = [
        ("heading 1",  NAVY,  16, True,  False),
        ("heading 2",  GOLD,  14, True,  False),
        ("heading 3",  DARK,  12, True,  True),
    ]
    for fragment, color, size, bold, italic in mappings:
        s = get_style(fragment)
        if s:
            s.font.name      = "Times New Roman"
            s.font.size      = Pt(size)
            s.font.bold      = bold
            s.font.italic    = italic
            s.font.color.rgb = RGBColor(*color)

    normal = get_style("Normal") or get_style("Default")
    if normal:
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)

    # ── Per-paragraph formatting ───────────────────────────────────────────────
    for para in doc.paragraphs:
        sn = para.style.name.lower()

        if any(x in sn for x in ("heading 1", "heading1")) or sn.endswith(" 1"):
            para.paragraph_format.space_before = Pt(18)
            para.paragraph_format.space_after  = Pt(8)
            for r in para.runs:
                r.font.name      = "Times New Roman"
                r.font.size      = Pt(16)
                r.font.bold      = True
                r.font.color.rgb = RGBColor(*NAVY)

        elif any(x in sn for x in ("heading 2", "heading2")) or sn.endswith(" 2"):
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after  = Pt(6)
            for r in para.runs:
                r.font.name      = "Times New Roman"
                r.font.size      = Pt(14)
                r.font.bold      = True
                r.font.color.rgb = RGBColor(*GOLD)

        elif any(x in sn for x in ("heading 3", "heading3")) or sn.endswith(" 3"):
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after  = Pt(4)
            for r in para.runs:
                r.font.name      = "Times New Roman"
                r.font.size      = Pt(12)
                r.font.bold      = True
                r.font.italic    = True
                r.font.color.rgb = RGBColor(*DARK)

        else:
            # Body / Normal
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.line_spacing = Pt(18)
            para.paragraph_format.space_after  = Pt(8)
            for r in para.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)

    # ── Tables ────────────────────────────────────────────────────────────────
    def shade_cell(cell, rgb):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "{:02X}{:02X}{:02X}".format(*rgb))
        tcPr.append(shd)

    for table in doc.tables:
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after  = Pt(3)
                    for r in p.runs:
                        r.font.name = "Times New Roman"
                        r.font.size = Pt(10)
                if i == 0:
                    shade_cell(cell, (27, 58, 107))
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.color.rgb = RGBColor(255, 255, 255)
                            r.font.bold      = True
                elif i % 2 == 0:
                    shade_cell(cell, LTBLU)

    # ── Header image ──────────────────────────────────────────────────────────
    if HEADER_IMG.exists():
        first = doc.paragraphs[0]
        new_p = OxmlElement("w:p")
        first._element.addprevious(new_p)
        img_para = doc.paragraphs[0]
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.paragraph_format.space_after = Pt(20)
        img_para.add_run().add_picture(str(HEADER_IMG), width=Inches(6.2))
        print("✓  Header image inserted")
    else:
        print("⚠  Header image not found at:", HEADER_IMG)

    # ── Footer with page numbers ───────────────────────────────────────────────
    def add_footer(section):
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        for ftype in ("begin", "end"):
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), ftype)
            if ftype == "begin":
                fc_begin = fc
            else:
                fc_end = fc
        instr = OxmlElement("w:instrText")
        instr.text = " PAGE "
        run._r.append(fc_begin)
        run._r.append(instr)
        run._r.append(fc_end)

    for sec in doc.sections:
        add_footer(sec)

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(str(FINAL_DOCX))
    print("✓  Saved →", FINAL_DOCX.name)
    print()
    print("═" * 55)
    print(f"  ✅  {FINAL_DOCX}")
    print("═" * 55)


# ─── MAIN ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    run_pandoc()
    apply_formatting()
    if RAW_DOCX.exists():
        RAW_DOCX.unlink()
        print("🗑  Removed intermediate file")
